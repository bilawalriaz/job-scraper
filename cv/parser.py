"""CV parsing module for extracting text and structure from DOCX files."""

import json
import logging
from pathlib import Path
from typing import Dict, Optional
from docx import Document

logger = logging.getLogger('cv.parser')


class CVParser:
    """Parse DOCX CV files and extract structured data."""

    def __init__(self, file_path: str):
        """Initialize parser with path to DOCX file."""
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"CV file not found: {file_path}")
        if not self.file_path.suffix.lower() == '.docx':
            raise ValueError("Only DOCX files are supported")
        self.doc = Document(file_path)

    def extract_text(self) -> str:
        """
        Extract all text from the document.

        Returns:
            Full text content of the document with paragraphs separated by newlines.
        """
        full_text = []

        # Extract from paragraphs
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        # Extract from tables
        for table in self.doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(' | '.join(row_text))

        return '\n'.join(full_text)

    def extract_sections(self) -> Dict[str, str]:
        """
        Extract document sections based on heading styles.

        Returns:
            Dictionary mapping section names to their content.
        """
        sections = {}
        current_section = 'header'
        current_content = []

        for para in self.doc.paragraphs:
            # Check if this is a heading
            style_name = para.style.name.lower() if para.style else ''
            text = para.text.strip()

            if 'heading' in style_name and text:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                # Start new section
                current_section = text.lower()
                current_content = []
            elif text:
                current_content.append(text)

        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)

        return sections

    def get_formatting_info(self) -> Dict:
        """
        Extract formatting information to preserve document structure.

        Returns:
            Dictionary with font, size, and style information.
        """
        formatting = {
            'fonts_used': set(),
            'has_tables': len(self.doc.tables) > 0,
            'section_count': len(self.doc.sections),
            'paragraph_count': len(self.doc.paragraphs)
        }

        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    formatting['fonts_used'].add(run.font.name)

        formatting['fonts_used'] = list(formatting['fonts_used'])
        return formatting


def parse_cv_with_llm(raw_text: str, llm_processor) -> Optional[Dict]:
    """
    Parse CV text using LLM to extract structured data.

    Args:
        raw_text: The raw text extracted from the CV
        llm_processor: The LLM processor instance

    Returns:
        Structured CV data or None if parsing fails
    """
    from llm.prompts import CV_PARSE_PROMPT

    prompt = CV_PARSE_PROMPT.format(cv_text=raw_text)

    try:
        result = llm_processor._call_llm_with_prompt(prompt)
        if result:
            return json.loads(result)
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse CV LLM response as JSON: {e}")
    except Exception as e:
        logger.error(f"Error parsing CV with LLM: {e}")

    return None
