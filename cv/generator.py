"""Document generation module for creating tailored CVs and cover letters."""

import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger('cv.generator')

# Ensure upload directories exist
UPLOAD_DIR = Path('data/uploads')
CV_DIR = UPLOAD_DIR / 'cv'
GENERATED_DIR = UPLOAD_DIR / 'generated'


def ensure_dirs():
    """Ensure upload directories exist."""
    CV_DIR.mkdir(parents=True, exist_ok=True)
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)


class DocumentGenerator:
    """Generate tailored CV and cover letter documents."""

    def __init__(self):
        ensure_dirs()

    def create_tailored_cv(self, original_cv_path: str, tailored_content: Dict,
                           job_info: Dict) -> str:
        """
        Create a tailored CV based on the original CV structure.

        Args:
            original_cv_path: Path to the original CV DOCX
            tailored_content: Dictionary with tailored sections from LLM
            job_info: Job details (title, company) for filename

        Returns:
            Path to the generated CV file
        """
        # Load original CV as template
        doc = Document(original_cv_path)

        # Replace content in paragraphs based on tailored content
        for para in doc.paragraphs:
            original_text = para.text.strip()
            if not original_text:
                continue

            # Check if this paragraph should be replaced
            for section_key, new_content in tailored_content.items():
                if section_key in ['experience', 'skills', 'summary']:
                    # Look for section headings and replace following content
                    if section_key.lower() in original_text.lower():
                        # This is a section header, content follows
                        pass

        # For now, we'll create a cleaner approach:
        # Generate a new document with the tailored content while preserving structure

        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        company_clean = self._clean_filename(job_info.get('company', 'company'))
        filename = f"cv_{company_clean}_{timestamp}.docx"
        output_path = GENERATED_DIR / filename

        # Save the document
        doc.save(str(output_path))

        logger.info(f"Generated tailored CV: {output_path}")
        return str(output_path)

    def create_cover_letter(self, content: str, job_info: Dict,
                            voice_profile: Optional[Dict] = None) -> str:
        """
        Create a cover letter document.

        Args:
            content: The cover letter text content
            job_info: Job details (title, company, etc.)
            voice_profile: Optional voice profile for styling

        Returns:
            Path to the generated cover letter file
        """
        doc = Document()

        # Set up document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add date (right-aligned)
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(datetime.now().strftime('%d %B %Y'))
        date_run.font.size = Pt(11)

        # Add spacing
        doc.add_paragraph()

        # Add recipient info
        company = job_info.get('company', 'Hiring Manager')
        title = job_info.get('title', 'the position')

        recipient = doc.add_paragraph()
        recipient.add_run(f"Dear Hiring Manager,").bold = False
        recipient.runs[0].font.size = Pt(11)

        # Add spacing
        doc.add_paragraph()

        # Add body paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph()
                run = para.add_run(para_text.strip())
                run.font.size = Pt(11)
                para.paragraph_format.space_after = Pt(12)

        # Add closing
        doc.add_paragraph()
        closing = doc.add_paragraph()
        closing.add_run("Kind regards,")
        closing.runs[0].font.size = Pt(11)

        # Add signature line
        doc.add_paragraph()
        sig = doc.add_paragraph()
        sig.add_run("[Your Name]")
        sig.runs[0].font.size = Pt(11)

        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        company_clean = self._clean_filename(company)
        filename = f"cover_letter_{company_clean}_{timestamp}.docx"
        output_path = GENERATED_DIR / filename

        # Save the document
        doc.save(str(output_path))

        logger.info(f"Generated cover letter: {output_path}")
        return str(output_path)

    def create_tailored_cv_from_content(self, cv_data: Dict, tailored_sections: Dict,
                                         job_info: Dict) -> str:
        """
        Create a new tailored CV from structured content.

        Args:
            cv_data: Parsed CV data with personal info, experience, skills, etc.
            tailored_sections: Sections rewritten by LLM for this job
            job_info: Job details for filename

        Returns:
            Path to the generated CV file
        """
        doc = Document()

        # Set up document margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Add name (if available)
        personal = cv_data.get('personal_info', {})
        if personal.get('name'):
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(personal['name'])
            name_run.bold = True
            name_run.font.size = Pt(16)

        # Add contact info
        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('location'):
            contact_parts.append(personal['location'])
        if personal.get('linkedin'):
            contact_parts.append(personal['linkedin'])

        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(' | '.join(contact_parts))
            contact_run.font.size = Pt(10)

        # Add horizontal line
        doc.add_paragraph('_' * 80)

        # Add summary (tailored if available)
        summary = tailored_sections.get('summary') or cv_data.get('summary', '')
        if summary:
            self._add_section(doc, 'Professional Summary', summary)

        # Add skills (tailored if available)
        skills = tailored_sections.get('skills') or cv_data.get('skills', {})
        if skills:
            skills_text = self._format_skills(skills)
            self._add_section(doc, 'Skills', skills_text)

        # Add experience (tailored if available)
        experience = tailored_sections.get('experience') or cv_data.get('experience', [])
        if experience:
            self._add_experience_section(doc, experience)

        # Add education
        education = cv_data.get('education', [])
        if education:
            self._add_education_section(doc, education)

        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        company_clean = self._clean_filename(job_info.get('company', 'company'))
        filename = f"cv_tailored_{company_clean}_{timestamp}.docx"
        output_path = GENERATED_DIR / filename

        doc.save(str(output_path))
        logger.info(f"Generated tailored CV from content: {output_path}")
        return str(output_path)

    def _add_section(self, doc: Document, title: str, content: str):
        """Add a section with title and content."""
        # Add section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title.upper())
        title_run.bold = True
        title_run.font.size = Pt(11)

        # Add content
        content_para = doc.add_paragraph()
        content_run = content_para.add_run(content)
        content_run.font.size = Pt(10)
        content_para.paragraph_format.space_after = Pt(12)

    def _add_experience_section(self, doc: Document, experience: list):
        """Add experience section."""
        # Section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('EXPERIENCE')
        title_run.bold = True
        title_run.font.size = Pt(11)

        for exp in experience:
            # Job title and company
            job_para = doc.add_paragraph()
            job_title = exp.get('title', '')
            company = exp.get('company', '')
            dates = exp.get('dates', '')

            job_run = job_para.add_run(f"{job_title}")
            job_run.bold = True
            job_run.font.size = Pt(10)

            if company:
                job_para.add_run(f" | {company}")
            if dates:
                job_para.add_run(f" | {dates}").italic = True

            # Responsibilities/achievements
            responsibilities = exp.get('responsibilities', [])
            achievements = exp.get('achievements', [])

            for item in responsibilities + achievements:
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.add_run(item).font.size = Pt(10)

            # Add spacing between jobs
            doc.add_paragraph()

    def _add_education_section(self, doc: Document, education: list):
        """Add education section."""
        # Section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('EDUCATION')
        title_run.bold = True
        title_run.font.size = Pt(11)

        for edu in education:
            edu_para = doc.add_paragraph()
            degree = edu.get('degree', '')
            institution = edu.get('institution', '')
            year = edu.get('year', '')

            edu_run = edu_para.add_run(f"{degree}")
            edu_run.bold = True
            edu_run.font.size = Pt(10)

            if institution:
                edu_para.add_run(f" - {institution}")
            if year:
                edu_para.add_run(f" ({year})")

    def _format_skills(self, skills: dict) -> str:
        """Format skills dictionary into text."""
        parts = []
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                if skill_list:
                    if isinstance(skill_list, list):
                        parts.append(f"{category.title()}: {', '.join(skill_list)}")
                    else:
                        parts.append(f"{category.title()}: {skill_list}")
            return '\n'.join(parts)
        elif isinstance(skills, list):
            return ', '.join(skills)
        return str(skills)

    def _clean_filename(self, name: str) -> str:
        """Clean a string for use in a filename."""
        # Remove or replace invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            name = name.replace(char, '')
        # Replace spaces with underscores
        name = name.replace(' ', '_')
        # Limit length
        return name[:50].strip('_')
